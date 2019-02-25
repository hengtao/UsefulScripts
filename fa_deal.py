# -*- coding: utf-8 -*-
"""
Created on Thu Jan 17 09:42:44 2019

@author: Administrator
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.dml import MSO_THEME_COLOR
from docx.enum.text import WD_UNDERLINE, WD_COLOR_INDEX
import re, sys, glob, os


def judge(info_total, p, tag_info):
	if tag_info == "##UPSTREAM##":
		######color#######
		p.add_run(info_total).font.color.theme_color = MSO_THEME_COLOR.ACCENT_6
	elif tag_info == '##UTR5##':
		######double underline####
		p.add_run(info_total).font.underline = WD_UNDERLINE.DOUBLE
	elif tag_info == '##CDS##':
		#######color########
		info_cds = re.split(r'\s+', info_total)
		for s in range(1, len(info_cds), 2):
			p.add_run(info_cds[s - 1]).font.color.theme_color = MSO_THEME_COLOR.ACCENT_2
			p.add_run(info_cds[s]).font.color.theme_color = MSO_THEME_COLOR.ACCENT_1
		if len(info_cds) % 2 == 1: p.add_run(info_cds[-1]).font.color.theme_color = MSO_THEME_COLOR.ACCENT_2

	elif tag_info == '##INTRON##':
		#######same##########
		p.add_run(info_total)
	elif tag_info == '##UTR3##':
		#####single underline######
		p.add_run(info_total).font.underline = True
	elif tag_info == '##DOWNSTREAM##':
		#########color##########
		p.add_run(info_total).font.color.theme_color = MSO_THEME_COLOR.ACCENT_3


def gene_deal(gene_info, p):
	#########color##########
	# ACCENT_1:blue
	# ACCENT_2:dark red
	# ACCENT_3:green
	# ACCENT_4:purple
	# ACCENT_5:shadow blue
	# ACCENT_6:orange

	info_total = re.split(r'##[A-Z0-9]*##', gene_info)[1]
	tag_info = re.match(r'##[A-Z0-9]*##', gene_info).group(0)
	var_info = re.findall(r'\[.*?\]', info_total)
	if len(var_info) == 0:
		judge(info_total, p, tag_info)
	else:
		no_var_info = re.split(r'\[.*?\]', info_total)
		for i in range(len(var_info)):
			judge(no_var_info[i], p, tag_info)
			hlseq = var_info[i].replace(" ","")
			hlseq = re.sub(r"[\:]{1,}", ":", hlseq)
			p.add_run(hlseq).font.highlight_color = WD_COLOR_INDEX.YELLOW
		judge(no_var_info[len(var_info)], p, tag_info)
	p.add_run(' ')


def main():
	document = Document()
	#############sys.argv[1]输入的是源文件的目录##############
	for i in glob.glob('%s/*.fa' % sys.argv[1]):
		with open(i, 'r') as f:
			for line in f:
				if line.startswith('>'):
					document.add_heading(line.strip(), 0)
					p = document.add_paragraph('')
					continue
				elif not line.strip() or line.startswith('---'):
					continue
				gene_deal(line.strip(), p)

		#document.add_page_break()
		#########默认word文件生成在源文件所在的目录，如需更改，可以改i#########
		document.save(i + '.docx')
		document = Document()
		#########删除源文件 如需保留可以注释########
		#os.remove(i)
main()